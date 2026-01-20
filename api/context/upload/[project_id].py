from http.server import BaseHTTPRequestHandler
import json
import os
import uuid
import io
from supabase import create_client

# File processing libraries
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

import email
from email import policy

# Configuration
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'xlsx', 'md', 'eml', 'csv'}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB


def get_supabase():
    url = os.environ.get('SUPABASE_URL')
    key = os.environ.get('SUPABASE_KEY')
    if not url or not key:
        raise Exception("Supabase credentials not configured")
    return create_client(url, key)


def cors_headers():
    return {
        'Access-Control-Allow-Origin': '*',
        'Access-Control-Allow-Methods': 'GET, POST, OPTIONS',
        'Access-Control-Allow-Headers': 'Content-Type, Authorization',
    }


def get_project_id(path):
    """Extract project_id from path like /api/context/upload/[project_id]"""
    parts = path.strip('/').split('/')
    if len(parts) >= 4:
        return parts[3]
    return None


def validate_uuid(uuid_str):
    """Validate UUID format"""
    try:
        uuid.UUID(uuid_str)
        return True
    except (ValueError, AttributeError):
        return False


def get_file_extension(filename):
    """Get file extension from filename"""
    return filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''


def extract_text_plain(file_data):
    """Extract text from plain text files"""
    try:
        return file_data.decode('utf-8')
    except UnicodeDecodeError:
        return file_data.decode('latin-1')


def extract_pdf(file_data):
    """Extract text from PDF files"""
    if PdfReader is None:
        return "PDF extraction not available"
    pdf_file = io.BytesIO(file_data)
    reader = PdfReader(pdf_file)
    text_parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)
    return '\n\n'.join(text_parts)


def extract_docx(file_data):
    """Extract text from Word documents"""
    if Document is None:
        return "DOCX extraction not available"
    docx_file = io.BytesIO(file_data)
    doc = Document(docx_file)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(' | '.join(row_text))
    return '\n'.join(text_parts)


def extract_xlsx(file_data):
    """Extract text from Excel files"""
    if openpyxl is None:
        return "XLSX extraction not available"
    xlsx_file = io.BytesIO(file_data)
    workbook = openpyxl.load_workbook(xlsx_file, data_only=True)
    text_parts = []
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        text_parts.append(f"=== Sheet: {sheet_name} ===")
        for row in sheet.iter_rows():
            row_values = []
            for cell in row:
                if cell.value is not None:
                    row_values.append(str(cell.value))
            if row_values:
                text_parts.append(' | '.join(row_values))
    return '\n'.join(text_parts)


def extract_email(file_data):
    """Extract text from email files"""
    msg = email.message_from_bytes(file_data, policy=policy.default)
    text_parts = []
    text_parts.append(f"From: {msg.get('From', 'N/A')}")
    text_parts.append(f"To: {msg.get('To', 'N/A')}")
    text_parts.append(f"Subject: {msg.get('Subject', 'N/A')}")
    text_parts.append(f"Date: {msg.get('Date', 'N/A')}")
    text_parts.append("")

    if msg.is_multipart():
        for part in msg.walk():
            content_type = part.get_content_type()
            if content_type == 'text/plain':
                payload = part.get_payload(decode=True)
                if payload:
                    try:
                        text_parts.append(payload.decode('utf-8'))
                    except UnicodeDecodeError:
                        text_parts.append(payload.decode('latin-1'))
    else:
        payload = msg.get_payload(decode=True)
        if payload:
            try:
                text_parts.append(payload.decode('utf-8'))
            except UnicodeDecodeError:
                text_parts.append(payload.decode('latin-1'))

    return '\n'.join(text_parts)


def extract_text(file_data, filename):
    """Extract text from file based on extension"""
    ext = get_file_extension(filename)

    extractors = {
        'txt': extract_text_plain,
        'md': extract_text_plain,
        'csv': extract_text_plain,
        'pdf': extract_pdf,
        'docx': extract_docx,
        'xlsx': extract_xlsx,
        'eml': extract_email
    }

    extractor = extractors.get(ext)
    if extractor:
        try:
            return extractor(file_data)
        except Exception as e:
            return f"Error extracting text: {str(e)}"
    return f"Unsupported file type: {ext}"


def parse_multipart(body, content_type):
    """Parse multipart/form-data manually"""
    files = []

    # Extract boundary from content-type
    boundary = None
    for part in content_type.split(';'):
        part = part.strip()
        if part.startswith('boundary='):
            boundary = part[9:].strip('"')
            break

    if not boundary:
        return files

    # Split body by boundary
    boundary_bytes = f'--{boundary}'.encode()
    parts = body.split(boundary_bytes)

    for part in parts:
        if not part or part == b'--\r\n' or part == b'--':
            continue

        # Split headers from content
        if b'\r\n\r\n' in part:
            header_section, content = part.split(b'\r\n\r\n', 1)
        elif b'\n\n' in part:
            header_section, content = part.split(b'\n\n', 1)
        else:
            continue

        # Parse headers
        headers = {}
        header_lines = header_section.decode('utf-8', errors='ignore').split('\n')
        for line in header_lines:
            if ':' in line:
                key, value = line.split(':', 1)
                headers[key.strip().lower()] = value.strip()

        # Check if this is a file field
        content_disp = headers.get('content-disposition', '')
        if 'filename=' in content_disp:
            # Extract filename
            filename = None
            for part_cd in content_disp.split(';'):
                part_cd = part_cd.strip()
                if part_cd.startswith('filename='):
                    filename = part_cd[9:].strip('"')
                    break

            if filename:
                # Remove trailing boundary markers and whitespace
                if content.endswith(b'\r\n'):
                    content = content[:-2]
                elif content.endswith(b'\n'):
                    content = content[:-1]

                files.append({
                    'filename': filename,
                    'data': content
                })

    return files


class handler(BaseHTTPRequestHandler):
    def send_cors_headers(self):
        for key, value in cors_headers().items():
            self.send_header(key, value)

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_cors_headers()
        self.end_headers()
        return

    def do_POST(self):
        """Upload context files for a project"""
        try:
            project_id = get_project_id(self.path)

            # Validate project_id
            if not project_id or not validate_uuid(project_id):
                self.send_response(400)
                self.send_cors_headers()
                self.send_header('Content-Type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'error': 'Invalid project ID format'}).encode())
                return

            # Check if project exists
            supabase = get_supabase()
            project_result = supabase.table('projects').select('*').eq('id', project_id).execute()
            if not project_result.data:
                self.send_response(404)
                self.send_cors_headers()
                self.send_header('Content-Type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'error': 'Project not found'}).encode())
                return

            # Read request body
            content_length = int(self.headers.get('Content-Length', 0))
            if content_length == 0:
                self.send_response(400)
                self.send_cors_headers()
                self.send_header('Content-Type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'error': 'No files provided'}).encode())
                return

            body = self.rfile.read(content_length)
            content_type = self.headers.get('Content-Type', '')

            # Parse multipart form data
            files = parse_multipart(body, content_type)

            if not files:
                self.send_response(400)
                self.send_cors_headers()
                self.send_header('Content-Type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'error': 'No valid files provided'}).encode())
                return

            uploaded = []
            errors = []

            for file_info in files:
                filename = file_info['filename']
                file_data = file_info['data']

                # Check file extension
                ext = get_file_extension(filename)
                if ext not in ALLOWED_EXTENSIONS:
                    errors.append({
                        'file': filename,
                        'error': f"Unsupported file type: .{ext}. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
                    })
                    continue

                # Check file size
                if len(file_data) > MAX_FILE_SIZE:
                    errors.append({
                        'file': filename,
                        'error': f"File too large. Maximum size: {MAX_FILE_SIZE // (1024*1024)}MB"
                    })
                    continue

                if len(file_data) == 0:
                    errors.append({
                        'file': filename,
                        'error': 'File is empty'
                    })
                    continue

                try:
                    # Extract text from file
                    extracted_text = extract_text(file_data, filename)

                    if extracted_text.startswith('Error extracting text:'):
                        errors.append({
                            'file': filename,
                            'error': extracted_text
                        })
                        continue

                    # Generate unique filename for storage
                    unique_filename = f"{project_id}/{uuid.uuid4()}_{filename}"

                    # Upload to Supabase Storage
                    file_url = ''
                    try:
                        storage_result = supabase.storage.from_('context-files').upload(
                            unique_filename,
                            file_data,
                            {'content-type': 'application/octet-stream'}
                        )
                        if storage_result:
                            file_url = supabase.storage.from_('context-files').get_public_url(unique_filename)
                    except Exception as storage_error:
                        print(f"Storage upload error: {storage_error}")
                        # Continue even if storage fails - we still have the extracted text

                    # Save metadata to database
                    file_id = str(uuid.uuid4())
                    db_result = supabase.table('context_files').insert({
                        'id': file_id,
                        'project_id': project_id,
                        'file_name': filename,
                        'file_type': ext,
                        'file_url': file_url,
                        'extracted_text': extracted_text
                    }).execute()

                    if db_result.data:
                        uploaded.append({
                            'id': file_id,
                            'file_name': filename,
                            'file_type': ext,
                            'text_length': len(extracted_text),
                            'text_preview': extracted_text[:500] + '...' if len(extracted_text) > 500 else extracted_text
                        })
                    else:
                        errors.append({
                            'file': filename,
                            'error': 'Failed to save file metadata to database'
                        })

                except Exception as e:
                    print(f"Error processing file {filename}: {e}")
                    errors.append({
                        'file': filename,
                        'error': f"Processing error: {str(e)}"
                    })

            self.send_response(200)
            self.send_cors_headers()
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({
                'uploaded': uploaded,
                'errors': errors,
                'summary': {
                    'total_files': len(files),
                    'successful': len(uploaded),
                    'failed': len(errors)
                }
            }).encode())

        except Exception as e:
            print(f"Error in upload_files: {e}")
            self.send_response(500)
            self.send_cors_headers()
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({'error': f'Upload failed: {str(e)}'}).encode())
        return
