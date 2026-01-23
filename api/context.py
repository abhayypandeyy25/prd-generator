from http.server import BaseHTTPRequestHandler
import json
import os
import uuid
import io
from supabase import create_client

# File processing libraries
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

import email
from email import policy
import re
import anthropic

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'xlsx', 'md', 'eml', 'csv'}
MAX_FILE_SIZE = 50 * 1024 * 1024

# Coverage categories for context analysis
COVERAGE_CATEGORIES = {
    'user_personas': {
        'keywords': ['user', 'persona', 'customer', 'audience', 'demographic', 'target user', 'end user', 'stakeholder'],
        'weight': 15,
        'description': 'User personas and target audience'
    },
    'success_metrics': {
        'keywords': ['metric', 'kpi', 'success', 'measure', 'goal', 'objective', 'okr', 'target', 'benchmark'],
        'weight': 15,
        'description': 'Success metrics and KPIs'
    },
    'technical_requirements': {
        'keywords': ['technical', 'api', 'database', 'architecture', 'system', 'integration', 'infrastructure', 'technology'],
        'weight': 15,
        'description': 'Technical requirements and constraints'
    },
    'timeline': {
        'keywords': ['timeline', 'deadline', 'milestone', 'schedule', 'date', 'quarter', 'sprint', 'release'],
        'weight': 10,
        'description': 'Timeline and milestones'
    },
    'user_stories': {
        'keywords': ['user story', 'use case', 'scenario', 'workflow', 'journey', 'as a user', 'i want to'],
        'weight': 15,
        'description': 'User stories and use cases'
    },
    'competitors': {
        'keywords': ['competitor', 'competition', 'market', 'alternative', 'benchmark', 'comparison', 'landscape'],
        'weight': 10,
        'description': 'Competitive analysis'
    },
    'constraints': {
        'keywords': ['constraint', 'limitation', 'budget', 'resource', 'dependency', 'risk', 'assumption'],
        'weight': 10,
        'description': 'Constraints and dependencies'
    },
    'scope': {
        'keywords': ['scope', 'feature', 'requirement', 'functionality', 'capability', 'in scope', 'out of scope'],
        'weight': 10,
        'description': 'Scope definition'
    }
}


def analyze_context_quality(files_data):
    """Analyze the quality and coverage of context files"""
    if not files_data:
        return {
            'quality_score': 0,
            'coverage': {},
            'suggestions': ['Upload context documents to improve PRD quality'],
            'summary': 'No context files uploaded'
        }

    # Combine all text
    all_text = '\n'.join([f.get('extracted_text', '') for f in files_data if f.get('extracted_text')])
    all_text_lower = all_text.lower()

    # Calculate base metrics
    total_length = len(all_text)
    file_count = len(files_data)
    unique_types = len(set(f.get('file_type', '') for f in files_data))

    # Length score (0-25 points)
    if total_length < 500:
        length_score = 5
    elif total_length < 2000:
        length_score = 10
    elif total_length < 5000:
        length_score = 15
    elif total_length < 15000:
        length_score = 20
    else:
        length_score = 25

    # Diversity score (0-15 points)
    diversity_score = min(15, file_count * 3 + unique_types * 2)

    # Coverage analysis (0-60 points)
    coverage = {}
    coverage_score = 0
    missing_categories = []

    for category, config in COVERAGE_CATEGORIES.items():
        matches = sum(1 for kw in config['keywords'] if kw in all_text_lower)
        if matches > 0:
            coverage[category] = {
                'found': True,
                'match_count': matches,
                'description': config['description']
            }
            coverage_score += config['weight']
        else:
            coverage[category] = {
                'found': False,
                'match_count': 0,
                'description': config['description']
            }
            missing_categories.append(config['description'])

    # Calculate total score
    total_score = min(100, length_score + diversity_score + coverage_score)

    # Generate suggestions
    suggestions = []
    if length_score < 15:
        suggestions.append('Add more detailed documentation for richer context')
    if diversity_score < 10:
        suggestions.append('Upload different types of documents (specs, research, emails) for diverse perspectives')
    if missing_categories:
        for cat in missing_categories[:3]:  # Top 3 missing
            suggestions.append(f'Add information about: {cat}')

    # Generate summary
    if total_score >= 80:
        summary = 'Excellent context coverage - ready for high-quality PRD generation'
    elif total_score >= 60:
        summary = 'Good context coverage with some gaps - PRD will be solid'
    elif total_score >= 40:
        summary = 'Moderate context - consider adding more documentation'
    else:
        summary = 'Limited context - PRD may have significant gaps'

    return {
        'quality_score': total_score,
        'coverage': coverage,
        'suggestions': suggestions,
        'summary': summary,
        'metrics': {
            'total_characters': total_length,
            'file_count': file_count,
            'unique_file_types': unique_types,
            'length_score': length_score,
            'diversity_score': diversity_score,
            'coverage_score': coverage_score
        }
    }


def extract_entities(text):
    """Extract key entities from text"""
    entities = {
        'dates': [],
        'percentages': [],
        'monetary': [],
        'technical_terms': []
    }

    # Date patterns
    date_patterns = [
        r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
        r'\b(?:Q[1-4]|H[1-2])\s*\d{4}\b',
        r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b',
        r'\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b'
    ]
    for pattern in date_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        entities['dates'].extend(matches[:10])  # Limit to 10

    # Percentages
    pct_matches = re.findall(r'\b\d+(?:\.\d+)?%\b', text)
    entities['percentages'] = list(set(pct_matches))[:10]

    # Monetary values
    money_matches = re.findall(r'\$[\d,]+(?:\.\d{2})?(?:\s*(?:million|billion|M|B|K))?\b', text, re.IGNORECASE)
    entities['monetary'] = list(set(money_matches))[:10]

    # Technical terms (common API/tech patterns)
    tech_patterns = [
        r'\b[A-Z]{2,}(?:\s+[A-Z]{2,})*\b',  # Acronyms
        r'\b(?:API|SDK|REST|GraphQL|OAuth|JWT|SSL|HTTP|HTTPS)\b'
    ]
    for pattern in tech_patterns:
        matches = re.findall(pattern, text)
        entities['technical_terms'].extend(matches[:20])
    entities['technical_terms'] = list(set(entities['technical_terms']))[:15]

    return entities


def detect_conflicts(text):
    """Detect potential conflicts or inconsistencies in the context"""
    conflicts = []
    text_lower = text.lower()

    # Check for contradictory statements
    contradiction_pairs = [
        ('must have', 'nice to have'),
        ('required', 'optional'),
        ('in scope', 'out of scope'),
        ('priority', 'deprioritize'),
        ('asap', 'later'),
        ('critical', 'low priority')
    ]

    for term1, term2 in contradiction_pairs:
        if term1 in text_lower and term2 in text_lower:
            conflicts.append({
                'type': 'potential_contradiction',
                'terms': [term1, term2],
                'message': f'Found both "{term1}" and "{term2}" - may need clarification'
            })

    return conflicts


def ai_analyze_context(text, project_description=''):
    """Use AI to provide deeper analysis of context"""
    api_key = os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        return None

    try:
        client = anthropic.Anthropic(api_key=api_key)

        prompt = f"""Analyze the following context documents for a product requirements document (PRD).
Project description: {project_description or 'Not provided'}

Context content (first 15000 chars):
{text[:15000]}

Provide a JSON response with:
1. "key_themes": List of 3-5 main themes/topics found
2. "stakeholders_mentioned": List of stakeholders/user types mentioned
3. "potential_risks": List of 2-3 potential risks or concerns
4. "missing_info": List of 2-3 critical missing pieces of information
5. "clarity_issues": List of any ambiguous or unclear statements
6. "recommendations": List of 2-3 actionable recommendations

Respond ONLY with valid JSON, no other text."""

        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1000,
            messages=[{"role": "user", "content": prompt}]
        )

        response_text = response.content[0].text.strip()
        # Try to parse JSON from response
        if response_text.startswith('{'):
            return json.loads(response_text)
        # Try to extract JSON from markdown code block
        json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', response_text, re.DOTALL)
        if json_match:
            return json.loads(json_match.group(1))
        return None
    except Exception as e:
        print(f"AI analysis error: {e}")
        return None


def summarize_file(text, filename):
    """Generate a summary of a single file using AI"""
    api_key = os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        return {'summary': text[:500] + '...' if len(text) > 500 else text, 'key_points': []}

    try:
        client = anthropic.Anthropic(api_key=api_key)

        prompt = f"""Summarize this document for a PRD context. File: {filename}

Content (first 10000 chars):
{text[:10000]}

Provide a JSON response with:
1. "summary": A 2-3 sentence summary
2. "key_points": List of 3-5 key points/facts
3. "document_type": Type of document (e.g., "user research", "technical spec", "meeting notes")
4. "relevance_score": 1-10 score of how relevant this is to product requirements

Respond ONLY with valid JSON."""

        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=500,
            messages=[{"role": "user", "content": prompt}]
        )

        response_text = response.content[0].text.strip()
        if response_text.startswith('{'):
            return json.loads(response_text)
        json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', response_text, re.DOTALL)
        if json_match:
            return json.loads(json_match.group(1))
        return {'summary': text[:500] + '...' if len(text) > 500 else text, 'key_points': []}
    except Exception:
        return {'summary': text[:500] + '...' if len(text) > 500 else text, 'key_points': []}


def get_supabase():
    url = os.environ.get('SUPABASE_URL')
    key = os.environ.get('SUPABASE_KEY')
    if not url or not key:
        raise Exception("Supabase credentials not configured")
    return create_client(url, key)


def cors_headers():
    return {
        'Access-Control-Allow-Origin': '*',
        'Access-Control-Allow-Methods': 'GET, POST, DELETE, OPTIONS',
        'Access-Control-Allow-Headers': 'Content-Type, Authorization',
        'Content-Type': 'application/json'
    }


def validate_uuid(uuid_str):
    try:
        uuid.UUID(uuid_str)
        return True
    except (ValueError, AttributeError):
        return False


def parse_path(path):
    """Parse path to determine operation type and IDs"""
    parts = path.strip('/').split('/')
    if len(parts) >= 4 and parts[2] == 'upload':
        return ('upload', parts[3], None)
    elif len(parts) >= 4 and parts[2] == 'text':
        return ('text', parts[3], None)
    elif len(parts) >= 4 and parts[2] == 'file':
        return ('file', None, parts[3])
    elif len(parts) >= 4 and parts[2] == 'analyze':
        return ('analyze', parts[3], None)
    elif len(parts) >= 4 and parts[2] == 'summarize':
        return ('summarize', None, parts[3])
    elif len(parts) >= 3:
        return ('list', parts[2], None)
    return (None, None, None)


def get_file_extension(filename):
    return filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''


def extract_text_plain(file_data):
    try:
        return file_data.decode('utf-8')
    except UnicodeDecodeError:
        return file_data.decode('latin-1')


def extract_pdf(file_data):
    if PdfReader is None:
        return "PDF extraction not available"
    pdf_file = io.BytesIO(file_data)
    reader = PdfReader(pdf_file)
    text_parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)
    return '\n\n'.join(text_parts)


def extract_docx(file_data):
    if Document is None:
        return "DOCX extraction not available"
    docx_file = io.BytesIO(file_data)
    doc = Document(docx_file)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(' | '.join(row_text))
    return '\n'.join(text_parts)


def extract_xlsx(file_data):
    if openpyxl is None:
        return "XLSX extraction not available"
    xlsx_file = io.BytesIO(file_data)
    workbook = openpyxl.load_workbook(xlsx_file, data_only=True)
    text_parts = []
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        text_parts.append(f"=== Sheet: {sheet_name} ===")
        for row in sheet.iter_rows():
            row_values = [str(cell.value) for cell in row if cell.value is not None]
            if row_values:
                text_parts.append(' | '.join(row_values))
    return '\n'.join(text_parts)


def extract_email(file_data):
    msg = email.message_from_bytes(file_data, policy=policy.default)
    text_parts = [
        f"From: {msg.get('From', 'N/A')}",
        f"To: {msg.get('To', 'N/A')}",
        f"Subject: {msg.get('Subject', 'N/A')}",
        f"Date: {msg.get('Date', 'N/A')}",
        ""
    ]
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == 'text/plain':
                payload = part.get_payload(decode=True)
                if payload:
                    try:
                        text_parts.append(payload.decode('utf-8'))
                    except UnicodeDecodeError:
                        text_parts.append(payload.decode('latin-1'))
    else:
        payload = msg.get_payload(decode=True)
        if payload:
            try:
                text_parts.append(payload.decode('utf-8'))
            except UnicodeDecodeError:
                text_parts.append(payload.decode('latin-1'))
    return '\n'.join(text_parts)


def extract_text(file_data, filename):
    ext = get_file_extension(filename)
    extractors = {
        'txt': extract_text_plain, 'md': extract_text_plain, 'csv': extract_text_plain,
        'pdf': extract_pdf, 'docx': extract_docx, 'xlsx': extract_xlsx, 'eml': extract_email
    }
    extractor = extractors.get(ext)
    if extractor:
        try:
            return extractor(file_data)
        except Exception as e:
            return f"Error extracting text: {str(e)}"
    return f"Unsupported file type: {ext}"


def parse_multipart(body, content_type):
    files = []
    boundary = None
    for part in content_type.split(';'):
        part = part.strip()
        if part.startswith('boundary='):
            boundary = part[9:].strip('"')
            break
    if not boundary:
        return files

    boundary_bytes = f'--{boundary}'.encode()
    parts = body.split(boundary_bytes)

    for part in parts:
        if not part or part == b'--\r\n' or part == b'--':
            continue
        if b'\r\n\r\n' in part:
            header_section, content = part.split(b'\r\n\r\n', 1)
        elif b'\n\n' in part:
            header_section, content = part.split(b'\n\n', 1)
        else:
            continue

        headers = {}
        for line in header_section.decode('utf-8', errors='ignore').split('\n'):
            if ':' in line:
                key, value = line.split(':', 1)
                headers[key.strip().lower()] = value.strip()

        content_disp = headers.get('content-disposition', '')
        if 'filename=' in content_disp:
            filename = None
            for part_cd in content_disp.split(';'):
                part_cd = part_cd.strip()
                if part_cd.startswith('filename='):
                    filename = part_cd[9:].strip('"')
                    break
            if filename:
                if content.endswith(b'\r\n'):
                    content = content[:-2]
                elif content.endswith(b'\n'):
                    content = content[:-1]
                files.append({'filename': filename, 'data': content})
    return files


class handler(BaseHTTPRequestHandler):
    def send_cors_headers(self):
        for key, value in cors_headers().items():
            self.send_header(key, value)

    def send_json(self, status, data):
        self.send_response(status)
        self.send_cors_headers()
        self.end_headers()
        self.wfile.write(json.dumps(data).encode())

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_cors_headers()
        self.end_headers()
        return

    def do_GET(self):
        """Handle GET requests for context files"""
        try:
            op, project_id, file_id = parse_path(self.path)
            supabase = get_supabase()

            if op == 'list' and project_id:
                result = supabase.table('context_files').select('*').eq('project_id', project_id).order('created_at', desc=True).execute()
                self.send_json(200, result.data if result.data else [])

            elif op == 'text' and project_id:
                result = supabase.table('context_files').select('file_name, extracted_text').eq('project_id', project_id).execute()
                texts = []
                for f in (result.data or []):
                    if f.get('extracted_text'):
                        texts.append(f"=== {f.get('file_name', 'Unknown')} ===\n{f['extracted_text']}")
                aggregated = "\n\n".join(texts)
                self.send_json(200, {'text': aggregated, 'length': len(aggregated), 'has_content': bool(aggregated.strip())})

            elif op == 'file' and file_id:
                if not validate_uuid(file_id):
                    self.send_json(400, {'error': 'Invalid file ID format'})
                    return
                result = supabase.table('context_files').select('*').eq('id', file_id).execute()
                if result.data:
                    self.send_json(200, result.data[0])
                else:
                    self.send_json(404, {'error': 'File not found'})

            elif op == 'analyze' and project_id:
                if not validate_uuid(project_id):
                    self.send_json(400, {'error': 'Invalid project ID format'})
                    return

                # Get all context files for project
                result = supabase.table('context_files').select('*').eq('project_id', project_id).execute()
                files_data = result.data or []

                # Perform quality analysis
                analysis = analyze_context_quality(files_data)

                # Combine all text for entity extraction
                all_text = '\n'.join([f.get('extracted_text', '') for f in files_data if f.get('extracted_text')])

                # Extract entities
                entities = extract_entities(all_text) if all_text else {}

                # Detect conflicts
                conflicts = detect_conflicts(all_text) if all_text else []

                self.send_json(200, {
                    **analysis,
                    'entities': entities,
                    'conflicts': conflicts,
                    'file_count': len(files_data)
                })

            elif op == 'summarize' and file_id:
                if not validate_uuid(file_id):
                    self.send_json(400, {'error': 'Invalid file ID format'})
                    return

                result = supabase.table('context_files').select('*').eq('id', file_id).execute()
                if not result.data:
                    self.send_json(404, {'error': 'File not found'})
                    return

                file_info = result.data[0]
                text = file_info.get('extracted_text', '')
                filename = file_info.get('file_name', 'unknown')

                summary = summarize_file(text, filename)
                self.send_json(200, {
                    'file_id': file_id,
                    'file_name': filename,
                    **summary
                })

            else:
                self.send_json(400, {'error': 'Invalid request path'})

        except Exception as e:
            self.send_json(500, {'error': str(e)})
        return

    def do_POST(self):
        """Handle file uploads and AI analysis"""
        try:
            op, project_id, file_id = parse_path(self.path)

            # Handle AI deep analysis
            if op == 'analyze' and project_id:
                if not validate_uuid(project_id):
                    self.send_json(400, {'error': 'Invalid project ID format'})
                    return

                supabase = get_supabase()

                # Get project info
                project_result = supabase.table('projects').select('*').eq('id', project_id).execute()
                project_description = ''
                if project_result.data:
                    project_description = project_result.data[0].get('description', '')

                # Get all context files
                result = supabase.table('context_files').select('*').eq('project_id', project_id).execute()
                files_data = result.data or []

                if not files_data:
                    self.send_json(400, {'error': 'No context files to analyze'})
                    return

                # Combine all text
                all_text = '\n'.join([f.get('extracted_text', '') for f in files_data if f.get('extracted_text')])

                # Perform basic analysis first
                basic_analysis = analyze_context_quality(files_data)
                entities = extract_entities(all_text)
                conflicts = detect_conflicts(all_text)

                # Perform AI deep analysis
                ai_analysis = ai_analyze_context(all_text, project_description)

                self.send_json(200, {
                    **basic_analysis,
                    'entities': entities,
                    'conflicts': conflicts,
                    'ai_analysis': ai_analysis,
                    'file_count': len(files_data)
                })
                return

            if op != 'upload' or not project_id:
                self.send_json(400, {'error': 'Invalid upload path'})
                return

            if not validate_uuid(project_id):
                self.send_json(400, {'error': 'Invalid project ID format'})
                return

            supabase = get_supabase()

            content_length = int(self.headers.get('Content-Length', 0))
            if content_length == 0:
                self.send_json(400, {'error': 'No files provided'})
                return

            body = self.rfile.read(content_length)
            content_type = self.headers.get('Content-Type', '')
            files = parse_multipart(body, content_type)

            if not files:
                self.send_json(400, {'error': 'No valid files provided'})
                return

            uploaded, errors = [], []
            for file_info in files:
                filename = file_info['filename']
                file_data = file_info['data']
                ext = get_file_extension(filename)

                if ext not in ALLOWED_EXTENSIONS:
                    errors.append({'file': filename, 'error': f"Unsupported file type: .{ext}"})
                    continue
                if len(file_data) > MAX_FILE_SIZE:
                    errors.append({'file': filename, 'error': 'File too large'})
                    continue
                if len(file_data) == 0:
                    errors.append({'file': filename, 'error': 'File is empty'})
                    continue

                try:
                    extracted_text = extract_text(file_data, filename)
                    if extracted_text.startswith('Error extracting text:'):
                        errors.append({'file': filename, 'error': extracted_text})
                        continue

                    unique_filename = f"{project_id}/{uuid.uuid4()}_{filename}"
                    file_url = ''
                    try:
                        supabase.storage.from_('context-files').upload(unique_filename, file_data, {'content-type': 'application/octet-stream'})
                        file_url = supabase.storage.from_('context-files').get_public_url(unique_filename)
                    except Exception:
                        pass

                    file_id = str(uuid.uuid4())
                    db_result = supabase.table('context_files').insert({
                        'id': file_id, 'project_id': project_id, 'file_name': filename,
                        'file_type': ext, 'file_url': file_url, 'extracted_text': extracted_text
                    }).execute()

                    if db_result.data:
                        uploaded.append({'id': file_id, 'file_name': filename, 'file_type': ext, 'text_length': len(extracted_text)})
                    else:
                        errors.append({'file': filename, 'error': 'Failed to save to database'})
                except Exception as e:
                    errors.append({'file': filename, 'error': str(e)})

            self.send_json(200, {'uploaded': uploaded, 'errors': errors, 'summary': {'total_files': len(files), 'successful': len(uploaded), 'failed': len(errors)}})

        except Exception as e:
            self.send_json(500, {'error': f'Upload failed: {str(e)}'})
        return

    def do_DELETE(self):
        """Delete a context file"""
        try:
            op, _, file_id = parse_path(self.path)

            if op != 'file' or not file_id:
                self.send_json(400, {'error': 'Invalid delete path'})
                return

            if not validate_uuid(file_id):
                self.send_json(400, {'error': 'Invalid file ID format'})
                return

            supabase = get_supabase()

            file_result = supabase.table('context_files').select('*').eq('id', file_id).execute()
            if not file_result.data:
                self.send_json(404, {'error': 'File not found'})
                return

            file_info = file_result.data[0]
            if file_info.get('file_url') and 'context-files/' in file_info['file_url']:
                try:
                    storage_path = file_info['file_url'].split('context-files/')[-1]
                    supabase.storage.from_('context-files').remove([storage_path])
                except Exception:
                    pass

            supabase.table('context_files').delete().eq('id', file_id).execute()
            self.send_json(200, {'message': 'File deleted successfully'})

        except Exception as e:
            self.send_json(500, {'error': f'Delete failed: {str(e)}'})
        return
