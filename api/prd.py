from http.server import BaseHTTPRequestHandler
import json
import os
import uuid
import io
import re
from supabase import create_client

try:
    import anthropic
except ImportError:
    anthropic = None

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None


PRD_TEMPLATE = """# Product Requirements Document

## Executive Summary
[Brief overview of the product and its purpose]

## Problem Statement
[Detailed description of the problem being solved]

## Target Users
[Description of primary and secondary users]

## Proposed Solution
[High-level description of the solution]

## Key Features
[List of main features]

## Success Metrics
[How success will be measured]

## Technical Requirements
[Technical considerations and requirements]

## Timeline and Milestones
[Key milestones and timeline]

## Risks and Mitigations
[Potential risks and how to address them]
"""


def load_questions():
    questions_file = os.path.join(os.path.dirname(__file__), 'data', 'questions.json')
    try:
        with open(questions_file, 'r') as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return {"sections": []}


def get_supabase():
    url = os.environ.get('SUPABASE_URL')
    key = os.environ.get('SUPABASE_KEY')
    if not url or not key:
        raise Exception("Supabase credentials not configured")
    return create_client(url, key)


def cors_headers():
    return {
        'Access-Control-Allow-Origin': '*',
        'Access-Control-Allow-Methods': 'GET, POST, PUT, DELETE, OPTIONS',
        'Access-Control-Allow-Headers': 'Content-Type, Authorization',
        'Content-Type': 'application/json'
    }


def validate_uuid(uuid_str):
    try:
        uuid.UUID(uuid_str)
        return True
    except (ValueError, AttributeError):
        return False


def verify_project_exists(supabase, project_id):
    """Verify that the project exists"""
    result = supabase.table('projects').select('id').eq('id', project_id).execute()
    return len(result.data) > 0 if result.data else False


def parse_path(path):
    """Parse path to determine operation and extract project_id
    /api/prd/{project_id} -> ('get', project_id)
    /api/prd/generate/{project_id} -> ('generate', project_id)
    /api/prd/preview/{project_id} -> ('preview', project_id)
    /api/prd/export/md/{project_id} -> ('export_md', project_id)
    /api/prd/export/docx/{project_id} -> ('export_docx', project_id)
    /api/prd/edit/{project_id} -> ('edit', project_id)
    /api/prd/history/{project_id} -> ('history', project_id)
    /api/prd/restore/{project_id}/{snapshot_id} -> ('restore', project_id, snapshot_id)
    /api/prd/regenerate-section/{project_id} -> ('regenerate_section', project_id)
    /api/prd/save-version/{project_id} -> ('save_version', project_id)
    /api/prd/compare/{project_id} -> ('compare', project_id)
    /api/prd/changelog/{project_id} -> ('changelog', project_id)
    /api/prd/snapshot/{snapshot_id} -> ('get_snapshot', snapshot_id)
    """
    parts = path.strip('/').split('/')
    if len(parts) >= 3:
        if parts[2] == 'generate' and len(parts) >= 4:
            return ('generate', parts[3], None)
        elif parts[2] == 'preview' and len(parts) >= 4:
            return ('preview', parts[3], None)
        elif parts[2] == 'export' and len(parts) >= 5:
            export_type = parts[3]
            project_id = parts[4]
            return (f'export_{export_type}', project_id, None)
        elif parts[2] == 'edit' and len(parts) >= 4:
            return ('edit', parts[3], None)
        elif parts[2] == 'history' and len(parts) >= 4:
            return ('history', parts[3], None)
        elif parts[2] == 'restore' and len(parts) >= 5:
            return ('restore', parts[3], parts[4])
        elif parts[2] == 'regenerate-section' and len(parts) >= 4:
            return ('regenerate_section', parts[3], None)
        elif parts[2] == 'save-version' and len(parts) >= 4:
            return ('save_version', parts[3], None)
        elif parts[2] == 'compare' and len(parts) >= 4:
            return ('compare', parts[3], None)
        elif parts[2] == 'changelog' and len(parts) >= 4:
            return ('changelog', parts[3], None)
        elif parts[2] == 'snapshot' and len(parts) >= 4:
            return ('get_snapshot', parts[3], None)
        else:
            return ('get', parts[2], None)
    return (None, None, None)


def get_question_map(questions_data):
    question_map = {}
    for section in questions_data.get('sections', []):
        for subsection in section.get('subsections', []):
            for q in subsection.get('questions', []):
                question_map[q['id']] = {'question': q.get('question', ''), 'section': section['title'], 'subsection': subsection['title']}
    return question_map


def generate_prd_with_claude(organized_responses, template):
    if anthropic is None:
        raise Exception("Anthropic library not available")
    api_key = os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        raise Exception("Anthropic API key not configured")

    client = anthropic.Anthropic(api_key=api_key)
    responses_text = ""
    for section_name, responses in organized_responses.items():
        responses_text += f"\n## {section_name}\n"
        for resp in responses:
            if resp.get('response') and resp.get('response').strip():
                responses_text += f"Q: {resp['question']}\nA: {resp['response']}\n\n"

    prompt = f"""You are a professional product manager. Based on the following Q&A responses, generate a comprehensive Product Requirements Document (PRD).

RESPONSES:
{responses_text}

TEMPLATE STRUCTURE:
{template}

Generate a well-structured PRD that:
1. Synthesizes all the provided answers into coherent sections
2. Follows the template structure
3. Uses professional language appropriate for stakeholders
4. Includes specific details from the responses
5. Adds appropriate formatting (headers, bullet points, etc.)

Output the PRD in Markdown format."""

    message = client.messages.create(model="claude-sonnet-4-20250514", max_tokens=8192, messages=[{"role": "user", "content": prompt}])
    return message.content[0].text


def regenerate_prd_section(current_prd, section_name, organized_responses):
    """Regenerate a specific section of the PRD while keeping other sections intact"""
    if anthropic is None:
        raise Exception("Anthropic library not available")
    api_key = os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        raise Exception("Anthropic API key not configured")

    client = anthropic.Anthropic(api_key=api_key)

    # Build context from responses
    responses_text = ""
    for section, responses in organized_responses.items():
        responses_text += f"\n### {section}\n"
        for resp in responses:
            if resp.get('response') and resp.get('response').strip():
                responses_text += f"Q: {resp['question']}\nA: {resp['response']}\n\n"

    prompt = f"""You are a professional product manager. I need you to regenerate ONLY the "{section_name}" section of the following PRD.

CURRENT PRD:
{current_prd}

AVAILABLE Q&A DATA (use this to improve the section):
{responses_text}

INSTRUCTIONS:
1. Find the "{section_name}" section in the PRD (look for ## {section_name} or similar headers)
2. Regenerate ONLY that section with improved content based on the Q&A data
3. Keep ALL other sections EXACTLY as they are
4. Return the COMPLETE PRD with the regenerated section

Output the complete PRD in Markdown format."""

    message = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=8192,
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text


def compute_diff(old_content, new_content):
    """Compute a simple diff between two texts"""
    import difflib

    old_lines = old_content.split('\n') if old_content else []
    new_lines = new_content.split('\n') if new_content else []

    differ = difflib.unified_diff(old_lines, new_lines, lineterm='')
    diff_lines = list(differ)

    added = sum(1 for line in diff_lines if line.startswith('+') and not line.startswith('+++'))
    removed = sum(1 for line in diff_lines if line.startswith('-') and not line.startswith('---'))

    return {
        'diff': '\n'.join(diff_lines),
        'added_lines': added,
        'removed_lines': removed,
        'total_changes': added + removed
    }


def generate_changelog(old_content, new_content, version_name=None):
    """Generate a human-readable changelog between two versions"""
    diff_result = compute_diff(old_content, new_content)

    # Extract section changes
    old_sections = set(re.findall(r'^##\s+(.+)$', old_content or '', re.MULTILINE))
    new_sections = set(re.findall(r'^##\s+(.+)$', new_content or '', re.MULTILINE))

    added_sections = new_sections - old_sections
    removed_sections = old_sections - new_sections
    modified_sections = []

    # Check which sections were modified
    for section in new_sections & old_sections:
        # Extract section content from both versions
        old_section_match = re.search(rf'^##\s+{re.escape(section)}\s*\n(.*?)(?=^##|\Z)', old_content, re.MULTILINE | re.DOTALL)
        new_section_match = re.search(rf'^##\s+{re.escape(section)}\s*\n(.*?)(?=^##|\Z)', new_content, re.MULTILINE | re.DOTALL)

        if old_section_match and new_section_match:
            if old_section_match.group(1).strip() != new_section_match.group(1).strip():
                modified_sections.append(section)

    changelog = f"# Changelog"
    if version_name:
        changelog += f" - {version_name}"
    changelog += f"\n\n"

    changelog += f"**Summary:** {diff_result['added_lines']} lines added, {diff_result['removed_lines']} lines removed\n\n"

    if added_sections:
        changelog += "## Added Sections\n"
        for section in added_sections:
            changelog += f"- {section}\n"
        changelog += "\n"

    if modified_sections:
        changelog += "## Modified Sections\n"
        for section in modified_sections:
            changelog += f"- {section}\n"
        changelog += "\n"

    if removed_sections:
        changelog += "## Removed Sections\n"
        for section in removed_sections:
            changelog += f"- {section}\n"
        changelog += "\n"

    return {
        'changelog': changelog,
        'added_sections': list(added_sections),
        'modified_sections': modified_sections,
        'removed_sections': list(removed_sections),
        'stats': diff_result
    }


def markdown_to_html(markdown_text):
    if not markdown_text:
        return ''
    html = markdown_text
    html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
    html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
    html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
    html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
    html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)
    html = re.sub(r'^- (.+)$', r'<li>\1</li>', html, flags=re.MULTILINE)
    html = re.sub(r'```(\w*)\n(.*?)```', r'<pre><code>\2</code></pre>', html, flags=re.DOTALL)
    html = re.sub(r'`(.+?)`', r'<code>\1</code>', html)
    html = re.sub(r'\n\n', '</p><p>', html)
    html = f'<p>{html}</p>'
    html = html.replace('<p></p>', '')
    return html


def markdown_to_docx(markdown_text, title="Product Requirements Document"):
    if Document is None:
        raise Exception("python-docx library not available")
    doc = Document()
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in markdown_text.split('\n'):
        line = line.rstrip()
        if not line:
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('- ') or line.startswith('* '):
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', line[2:])
            doc.add_paragraph(text, style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            doc.add_paragraph(text, style='List Number')
        else:
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            if text.strip():
                doc.add_paragraph(text)

    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer.getvalue()


class handler(BaseHTTPRequestHandler):
    def send_cors_headers(self):
        for key, value in cors_headers().items():
            self.send_header(key, value)

    def send_json(self, status, data):
        self.send_response(status)
        self.send_cors_headers()
        self.end_headers()
        self.wfile.write(json.dumps(data).encode())

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_cors_headers()
        self.end_headers()
        return

    def do_GET(self):
        try:
            op, project_id, extra_id = parse_path(self.path)

            if not project_id or not validate_uuid(project_id):
                self.send_json(400, {'error': 'Invalid project ID format'})
                return

            supabase = get_supabase()

            # Verify project exists
            if not verify_project_exists(supabase, project_id):
                self.send_json(404, {'error': 'Project not found'})
                return

            result = supabase.table('generated_prds').select('*').eq('project_id', project_id).order('created_at', desc=True).limit(1).execute()

            if op == 'get':
                if result.data:
                    self.send_json(200, result.data[0])
                else:
                    self.send_json(404, {'error': 'No PRD found. Please generate one first.'})

            elif op == 'preview':
                if not result.data:
                    self.send_json(404, {'error': 'No PRD found.', 'markdown': '', 'html': ''})
                    return
                prd = result.data[0]
                content = prd.get('content_md', '')
                html = markdown_to_html(content) if content else ''
                self.send_json(200, {'markdown': content, 'html': html, 'created_at': prd.get('created_at'), 'prd_id': prd.get('id')})

            elif op == 'export_md':
                if not result.data:
                    self.send_json(404, {'error': 'No PRD found.'})
                    return
                prd = result.data[0]
                content = prd.get('content_md', '')
                if not content:
                    self.send_json(404, {'error': 'PRD content is empty'})
                    return
                project_result = supabase.table('projects').select('name').eq('id', project_id).execute()
                project_name = project_result.data[0].get('name', 'PRD') if project_result.data else 'PRD'
                safe_name = "".join(c for c in project_name if c.isalnum() or c in (' ', '-', '_')).strip()
                md_bytes = content.encode('utf-8')
                self.send_response(200)
                self.send_header('Content-Type', 'text/markdown')
                self.send_header('Content-Disposition', f'attachment; filename=PRD_{safe_name}.md')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.send_header('Content-Length', str(len(md_bytes)))
                self.end_headers()
                self.wfile.write(md_bytes)

            elif op == 'export_docx':
                if Document is None:
                    self.send_json(500, {'error': 'Word document export not available'})
                    return
                if not result.data:
                    self.send_json(404, {'error': 'No PRD found.'})
                    return
                prd = result.data[0]
                content = prd.get('content_md', '')
                if not content:
                    self.send_json(404, {'error': 'PRD content is empty'})
                    return
                project_result = supabase.table('projects').select('name').eq('id', project_id).execute()
                project_name = project_result.data[0].get('name', 'Product') if project_result.data else 'Product'
                safe_name = "".join(c for c in project_name if c.isalnum() or c in (' ', '-', '_')).strip()
                title = f"PRD - {project_name}"
                try:
                    docx_bytes = markdown_to_docx(content, title)
                except Exception as e:
                    self.send_json(500, {'error': f'Failed to generate Word document: {str(e)}'})
                    return
                self.send_response(200)
                self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                self.send_header('Content-Disposition', f'attachment; filename=PRD_{safe_name}.docx')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.send_header('Content-Length', str(len(docx_bytes)))
                self.end_headers()
                self.wfile.write(docx_bytes)

            elif op == 'history':
                # Get edit history/snapshots for the PRD
                if not result.data:
                    self.send_json(200, {'snapshots': [], 'message': 'No PRD found'})
                    return
                prd = result.data[0]
                prd_id = prd.get('id')

                snapshots_result = supabase.table('prd_edit_snapshots').select(
                    'id, version_name, is_major_version, change_summary, created_at'
                ).eq('prd_id', prd_id).order('created_at', desc=True).execute()

                snapshots = snapshots_result.data if snapshots_result.data else []

                self.send_json(200, {
                    'prd_id': prd_id,
                    'current_content': prd.get('content_md', ''),
                    'is_manually_edited': prd.get('is_manually_edited', False),
                    'last_edited_at': prd.get('last_edited_at'),
                    'snapshots': snapshots
                })

            elif op == 'get_snapshot':
                # Get a specific snapshot's content
                snapshot_id = project_id  # In this case, project_id is actually snapshot_id
                if not validate_uuid(snapshot_id):
                    self.send_json(400, {'error': 'Invalid snapshot ID'})
                    return

                snapshot_result = supabase.table('prd_edit_snapshots').select('*').eq('id', snapshot_id).execute()
                if not snapshot_result.data:
                    self.send_json(404, {'error': 'Snapshot not found'})
                    return

                self.send_json(200, snapshot_result.data[0])

            else:
                self.send_json(400, {'error': 'Invalid request path'})

        except Exception as e:
            self.send_json(500, {'error': str(e)})
        return

    def do_PUT(self):
        """Handle PUT requests for editing PRD"""
        try:
            op, project_id, extra_id = parse_path(self.path)

            if op != 'edit':
                self.send_json(400, {'error': 'Invalid request path'})
                return

            if not project_id or not validate_uuid(project_id):
                self.send_json(400, {'error': 'Invalid project ID format'})
                return

            content_length = int(self.headers.get('Content-Length', 0))
            body = json.loads(self.rfile.read(content_length)) if content_length > 0 else {}

            new_content = body.get('content_md', body.get('content', '')).strip()
            if not new_content:
                self.send_json(400, {'error': 'Content is required'})
                return

            description = body.get('description', 'Manual edit')
            create_snapshot = body.get('create_snapshot', True)

            supabase = get_supabase()

            # Verify project exists
            if not verify_project_exists(supabase, project_id):
                self.send_json(404, {'error': 'Project not found'})
                return

            # Get current PRD
            result = supabase.table('generated_prds').select('*').eq('project_id', project_id).order('created_at', desc=True).limit(1).execute()

            if not result.data:
                self.send_json(404, {'error': 'No PRD found to edit'})
                return

            prd = result.data[0]
            prd_id = prd['id']
            old_content = prd.get('content_md', '')

            # Create snapshot of current content before updating (for undo)
            if create_snapshot and old_content:
                snapshot_id = str(uuid.uuid4())
                supabase.table('prd_edit_snapshots').insert({
                    'id': snapshot_id,
                    'prd_id': prd_id,
                    'project_id': project_id,
                    'snapshot_content': old_content,
                    'change_summary': description,
                    'is_major_version': False
                }).execute()

            # Update the PRD content
            supabase.table('generated_prds').update({
                'content_md': new_content,
                'is_manually_edited': True
            }).eq('id', prd_id).execute()

            self.send_json(200, {
                'success': True,
                'message': 'PRD updated successfully',
                'prd_id': prd_id
            })

        except Exception as e:
            self.send_json(500, {'error': str(e)})
        return

    def do_POST(self):
        try:
            op, project_id, extra_id = parse_path(self.path)

            if op not in ['generate', 'restore', 'regenerate_section', 'save_version', 'compare', 'changelog']:
                self.send_json(400, {'error': 'Invalid request path'})
                return

            if not project_id or not validate_uuid(project_id):
                self.send_json(400, {'error': 'Invalid project ID format'})
                return

            supabase = get_supabase()

            # Verify project exists
            if not verify_project_exists(supabase, project_id):
                self.send_json(404, {'error': 'Project not found'})
                return

            if op == 'generate':
                responses_result = supabase.table('question_responses').select('*').eq('project_id', project_id).execute()
                responses = responses_result.data if responses_result.data else []

                if not responses:
                    self.send_json(400, {'error': 'No responses found. Please answer questions first.'})
                    return

                confirmed_responses = [r for r in responses if r.get('confirmed')]
                if not confirmed_responses:
                    self.send_json(400, {'error': 'No confirmed responses found. Please confirm at least some answers.'})
                    return

                questions_data = load_questions()
                question_map = get_question_map(questions_data)
                organized_responses = {}

                for resp in responses:
                    q_id = resp.get('question_id')
                    if q_id in question_map:
                        section_key = question_map[q_id]['section']
                        if section_key not in organized_responses:
                            organized_responses[section_key] = []
                        organized_responses[section_key].append({
                            'question_id': q_id,
                            'question': question_map[q_id]['question'],
                            'response': resp.get('response', ''),
                            'confirmed': resp.get('confirmed', False)
                        })

                try:
                    prd_content = generate_prd_with_claude(organized_responses, PRD_TEMPLATE)
                except Exception as e:
                    self.send_json(503, {'error': 'AI service temporarily unavailable', 'details': str(e)})
                    return

                if not prd_content or prd_content.startswith('# Error'):
                    self.send_json(500, {'error': 'Failed to generate PRD content'})
                    return

                prd_id = str(uuid.uuid4())
                try:
                    prd_result = supabase.table('generated_prds').insert({'id': prd_id, 'project_id': project_id, 'content_md': prd_content}).execute()
                    saved_prd = prd_result.data[0] if prd_result.data else None
                except Exception as e:
                    self.send_json(500, {'error': 'PRD generated but failed to save', 'content': prd_content, 'details': str(e)})
                    return

                self.send_json(200, {
                    'message': 'PRD generated successfully',
                    'prd_id': saved_prd['id'] if saved_prd else prd_id,
                    'content': prd_content,
                    'stats': {'total_responses': len(responses), 'confirmed_responses': len(confirmed_responses), 'sections_covered': len(organized_responses)}
                })
                return

            elif op == 'restore':
            # Restore PRD from a snapshot
            snapshot_id = extra_id
            if not snapshot_id or not validate_uuid(snapshot_id):
                self.send_json(400, {'error': 'Invalid snapshot ID'})
                return

            content_length = int(self.headers.get('Content-Length', 0))
            body = json.loads(self.rfile.read(content_length)) if content_length > 0 else {}
            create_backup = body.get('create_backup', True)

            # Get the snapshot
            snapshot_result = supabase.table('prd_edit_snapshots').select('*').eq('id', snapshot_id).execute()
            if not snapshot_result.data:
                self.send_json(404, {'error': 'Snapshot not found'})
                return

            snapshot = snapshot_result.data[0]
            snapshot_content = snapshot.get('snapshot_content', '')
            prd_id = snapshot.get('prd_id')

            # Get current PRD
            prd_result = supabase.table('generated_prds').select('*').eq('id', prd_id).execute()
            if not prd_result.data:
                self.send_json(404, {'error': 'PRD not found'})
                return

            prd = prd_result.data[0]
            current_content = prd.get('content_md', '')

            # Create backup of current content before restoring
            if create_backup and current_content:
                backup_id = str(uuid.uuid4())
                supabase.table('prd_edit_snapshots').insert({
                    'id': backup_id,
                    'prd_id': prd_id,
                    'project_id': project_id,
                    'snapshot_content': current_content,
                    'change_summary': 'Backup before restore',
                    'is_major_version': False
                }).execute()

            # Restore the snapshot content
            supabase.table('generated_prds').update({
                'content_md': snapshot_content,
                'is_manually_edited': True
            }).eq('id', prd_id).execute()

            self.send_json(200, {
                'success': True,
                'message': 'PRD restored successfully',
                'content': snapshot_content,
                'prd_id': prd_id
            })
            return

        elif op == 'save_version':
            # Save current PRD as a named version
            content_length = int(self.headers.get('Content-Length', 0))
            body = json.loads(self.rfile.read(content_length)) if content_length > 0 else {}

            version_name = body.get('version_name', '').strip()
            change_summary = body.get('change_summary', '').strip()

            if not version_name:
                self.send_json(400, {'error': 'Version name is required'})
                return

            # Get current PRD
            result = supabase.table('generated_prds').select('*').eq('project_id', project_id).order('created_at', desc=True).limit(1).execute()

            if not result.data:
                self.send_json(404, {'error': 'No PRD found'})
                return

            prd = result.data[0]
            prd_id = prd['id']
            current_content = prd.get('content_md', '')

            # Create named version snapshot
            version_id = str(uuid.uuid4())
            supabase.table('prd_edit_snapshots').insert({
                'id': version_id,
                'prd_id': prd_id,
                'project_id': project_id,
                'snapshot_content': current_content,
                'version_name': version_name,
                'change_summary': change_summary,
                'is_major_version': True
            }).execute()

            self.send_json(200, {
                'success': True,
                'message': f'Version "{version_name}" saved successfully',
                'version_id': version_id
            })
            return

        elif op == 'regenerate_section':
            # Regenerate a specific section of the PRD
            content_length = int(self.headers.get('Content-Length', 0))
            body = json.loads(self.rfile.read(content_length)) if content_length > 0 else {}

            section_name = body.get('section_name', '').strip()
            if not section_name:
                self.send_json(400, {'error': 'Section name is required'})
                return

            # Get current PRD
            result = supabase.table('generated_prds').select('*').eq('project_id', project_id).order('created_at', desc=True).limit(1).execute()

            if not result.data:
                self.send_json(404, {'error': 'No PRD found'})
                return

            prd = result.data[0]
            prd_id = prd['id']
            current_content = prd.get('content_md', '')

            # Get responses for context
            responses_result = supabase.table('question_responses').select('*').eq('project_id', project_id).execute()
            responses = responses_result.data if responses_result.data else []

            questions_data = load_questions()
            question_map = get_question_map(questions_data)

            # Organize responses
            organized_responses = {}
            for resp in responses:
                q_id = resp.get('question_id')
                if q_id in question_map:
                    section_key = question_map[q_id]['section']
                    if section_key not in organized_responses:
                        organized_responses[section_key] = []
                    organized_responses[section_key].append({
                        'question_id': q_id,
                        'question': question_map[q_id]['question'],
                        'response': resp.get('response', ''),
                        'confirmed': resp.get('confirmed', False)
                    })

            # Generate only the requested section
            try:
                regenerated_section = regenerate_prd_section(current_content, section_name, organized_responses)
            except Exception as e:
                self.send_json(503, {'error': 'AI service temporarily unavailable', 'details': str(e)})
                return

            # Create snapshot before updating
            snapshot_id = str(uuid.uuid4())
            supabase.table('prd_edit_snapshots').insert({
                'id': snapshot_id,
                'prd_id': prd_id,
                'project_id': project_id,
                'snapshot_content': current_content,
                'change_summary': f'Before regenerating {section_name}',
                'is_major_version': False
            }).execute()

            # Update PRD with new section
            supabase.table('generated_prds').update({
                'content_md': regenerated_section,
                'is_manually_edited': True
            }).eq('id', prd_id).execute()

            self.send_json(200, {
                'success': True,
                'message': f'Section "{section_name}" regenerated successfully',
                'content': regenerated_section,
                'prd_id': prd_id
            })
            return

        elif op == 'compare':
            # Compare two versions
            content_length = int(self.headers.get('Content-Length', 0))
            body = json.loads(self.rfile.read(content_length)) if content_length > 0 else {}

            version1_id = body.get('version1_id')
            version2_id = body.get('version2_id')  # 'current' for current version

            # Get current PRD
            result = supabase.table('generated_prds').select('*').eq('project_id', project_id).order('created_at', desc=True).limit(1).execute()

            if not result.data:
                self.send_json(404, {'error': 'No PRD found'})
                return

            prd = result.data[0]
            current_content = prd.get('content_md', '')

            # Get version 1 content
            if version1_id == 'current':
                content1 = current_content
                version1_name = 'Current Version'
            else:
                if not validate_uuid(version1_id):
                    self.send_json(400, {'error': 'Invalid version1 ID'})
                    return
                snap1 = supabase.table('prd_edit_snapshots').select('*').eq('id', version1_id).execute()
                if not snap1.data:
                    self.send_json(404, {'error': 'Version 1 not found'})
                    return
                content1 = snap1.data[0].get('snapshot_content', '')
                version1_name = snap1.data[0].get('version_name') or snap1.data[0].get('created_at')

            # Get version 2 content
            if version2_id == 'current':
                content2 = current_content
                version2_name = 'Current Version'
            else:
                if not validate_uuid(version2_id):
                    self.send_json(400, {'error': 'Invalid version2 ID'})
                    return
                snap2 = supabase.table('prd_edit_snapshots').select('*').eq('id', version2_id).execute()
                if not snap2.data:
                    self.send_json(404, {'error': 'Version 2 not found'})
                    return
                content2 = snap2.data[0].get('snapshot_content', '')
                version2_name = snap2.data[0].get('version_name') or snap2.data[0].get('created_at')

            # Compute diff
            diff_result = compute_diff(content1, content2)

            self.send_json(200, {
                'version1': {'id': version1_id, 'name': version1_name, 'content': content1},
                'version2': {'id': version2_id, 'name': version2_name, 'content': content2},
                'diff': diff_result['diff'],
                'stats': {
                    'added_lines': diff_result['added_lines'],
                    'removed_lines': diff_result['removed_lines'],
                    'total_changes': diff_result['total_changes']
                }
            })
            return

        elif op == 'changelog':
            # Generate changelog between versions
            content_length = int(self.headers.get('Content-Length', 0))
            body = json.loads(self.rfile.read(content_length)) if content_length > 0 else {}

            from_version_id = body.get('from_version_id')
            to_version_id = body.get('to_version_id', 'current')
            version_name = body.get('version_name')

            # Get current PRD
            result = supabase.table('generated_prds').select('*').eq('project_id', project_id).order('created_at', desc=True).limit(1).execute()

            if not result.data:
                self.send_json(404, {'error': 'No PRD found'})
                return

            prd = result.data[0]
            current_content = prd.get('content_md', '')

            # Get from content
            if from_version_id == 'current':
                from_content = current_content
            else:
                if not validate_uuid(from_version_id):
                    self.send_json(400, {'error': 'Invalid from_version_id'})
                    return
                snap = supabase.table('prd_edit_snapshots').select('snapshot_content').eq('id', from_version_id).execute()
                if not snap.data:
                    self.send_json(404, {'error': 'From version not found'})
                    return
                from_content = snap.data[0].get('snapshot_content', '')

            # Get to content
            if to_version_id == 'current':
                to_content = current_content
            else:
                if not validate_uuid(to_version_id):
                    self.send_json(400, {'error': 'Invalid to_version_id'})
                    return
                snap = supabase.table('prd_edit_snapshots').select('snapshot_content').eq('id', to_version_id).execute()
                if not snap.data:
                    self.send_json(404, {'error': 'To version not found'})
                    return
                to_content = snap.data[0].get('snapshot_content', '')

            # Generate changelog
            changelog_result = generate_changelog(from_content, to_content, version_name)

            self.send_json(200, changelog_result)
            return

        except Exception as e:
            self.send_json(500, {'error': f'Operation failed: {str(e)}'})
        return
