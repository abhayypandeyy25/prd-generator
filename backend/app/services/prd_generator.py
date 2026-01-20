from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
import markdown


class PRDGenerator:
    """Generate PRD in various formats"""

    def __init__(self):
        self.md = markdown.Markdown(extensions=['tables', 'fenced_code'])

    def markdown_to_html(self, md_content: str) -> str:
        """Convert markdown to HTML"""
        self.md.reset()
        return self.md.convert(md_content)

    def markdown_to_docx(self, md_content: str, title: str = "Product Requirements Document") -> bytes:
        """Convert markdown PRD to Word document"""
        doc = Document()

        # Set document title
        doc.add_heading(title, 0)

        # Process markdown content
        lines = md_content.split('\n')
        current_list = None
        in_code_block = False
        code_content = []

        for line in lines:
            # Handle code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # End of code block
                    if code_content:
                        para = doc.add_paragraph()
                        para.style = 'Quote'
                        para.add_run('\n'.join(code_content)).font.name = 'Courier New'
                    code_content = []
                in_code_block = not in_code_block
                continue

            if in_code_block:
                code_content.append(line)
                continue

            # Handle headings
            if line.startswith('######'):
                doc.add_heading(line[6:].strip(), level=6)
            elif line.startswith('#####'):
                doc.add_heading(line[5:].strip(), level=5)
            elif line.startswith('####'):
                doc.add_heading(line[4:].strip(), level=4)
            elif line.startswith('###'):
                doc.add_heading(line[3:].strip(), level=3)
            elif line.startswith('##'):
                doc.add_heading(line[2:].strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(line[1:].strip(), level=1)

            # Handle horizontal rules
            elif line.strip() in ['---', '***', '___']:
                doc.add_paragraph('_' * 50)

            # Handle bullet lists
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:]
                para = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(para, text)

            # Handle numbered lists
            elif re.match(r'^\d+\.\s', line.strip()):
                text = re.sub(r'^\d+\.\s', '', line.strip())
                para = doc.add_paragraph(style='List Number')
                self._add_formatted_text(para, text)

            # Handle checkbox lists
            elif line.strip().startswith('- [ ]') or line.strip().startswith('- [x]'):
                checked = '[x]' in line[:10].lower()
                text = line.strip()[6:]
                para = doc.add_paragraph(style='List Bullet')
                checkbox = '\u2611' if checked else '\u2610'
                para.add_run(f"{checkbox} ")
                self._add_formatted_text(para, text)

            # Handle blockquotes
            elif line.strip().startswith('>'):
                text = line.strip()[1:].strip()
                para = doc.add_paragraph()
                para.style = 'Quote'
                self._add_formatted_text(para, text)

            # Handle tables (basic support)
            elif '|' in line and not line.strip().startswith('|---'):
                cells = [c.strip() for c in line.split('|') if c.strip()]
                if cells:
                    para = doc.add_paragraph()
                    para.add_run(' | '.join(cells))

            # Regular paragraphs
            elif line.strip():
                para = doc.add_paragraph()
                self._add_formatted_text(para, line)

        # Save to bytes
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io.getvalue()

    def _add_formatted_text(self, paragraph, text: str):
        """Add text with basic markdown formatting (bold, italic)"""
        # Simple parsing for bold and italic
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                # Italic
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # Code
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
            else:
                paragraph.add_run(part)

    def generate_markdown_file(self, md_content: str) -> bytes:
        """Return markdown content as bytes for download"""
        return md_content.encode('utf-8')


prd_generator = PRDGenerator()
