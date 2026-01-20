import io
import os
from PyPDF2 import PdfReader
from docx import Document
import openpyxl
import email
from email import policy


class FileProcessor:
    """Extract text from various file types"""

    SUPPORTED_TYPES = {
        'txt': 'text/plain',
        'md': 'text/markdown',
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'csv': 'text/csv',
        'eml': 'message/rfc822'
    }

    @staticmethod
    def get_file_extension(filename: str) -> str:
        return filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

    @staticmethod
    def extract_text(file_data: bytes, filename: str) -> str:
        """Extract text from file based on extension"""
        ext = FileProcessor.get_file_extension(filename)

        extractors = {
            'txt': FileProcessor._extract_text_plain,
            'md': FileProcessor._extract_text_plain,
            'csv': FileProcessor._extract_text_plain,
            'pdf': FileProcessor._extract_pdf,
            'docx': FileProcessor._extract_docx,
            'xlsx': FileProcessor._extract_xlsx,
            'eml': FileProcessor._extract_email
        }

        extractor = extractors.get(ext)
        if extractor:
            try:
                return extractor(file_data)
            except Exception as e:
                return f"Error extracting text: {str(e)}"
        return f"Unsupported file type: {ext}"

    @staticmethod
    def _extract_text_plain(file_data: bytes) -> str:
        """Extract text from plain text files"""
        try:
            return file_data.decode('utf-8')
        except UnicodeDecodeError:
            return file_data.decode('latin-1')

    @staticmethod
    def _extract_pdf(file_data: bytes) -> str:
        """Extract text from PDF files"""
        pdf_file = io.BytesIO(file_data)
        reader = PdfReader(pdf_file)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return '\n\n'.join(text_parts)

    @staticmethod
    def _extract_docx(file_data: bytes) -> str:
        """Extract text from Word documents"""
        docx_file = io.BytesIO(file_data)
        doc = Document(docx_file)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(' | '.join(row_text))
        return '\n'.join(text_parts)

    @staticmethod
    def _extract_xlsx(file_data: bytes) -> str:
        """Extract text from Excel files"""
        xlsx_file = io.BytesIO(file_data)
        workbook = openpyxl.load_workbook(xlsx_file, data_only=True)
        text_parts = []
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"=== Sheet: {sheet_name} ===")
            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value))
                if row_values:
                    text_parts.append(' | '.join(row_values))
        return '\n'.join(text_parts)

    @staticmethod
    def _extract_email(file_data: bytes) -> str:
        """Extract text from email files"""
        msg = email.message_from_bytes(file_data, policy=policy.default)
        text_parts = []

        # Headers
        text_parts.append(f"From: {msg.get('From', 'N/A')}")
        text_parts.append(f"To: {msg.get('To', 'N/A')}")
        text_parts.append(f"Subject: {msg.get('Subject', 'N/A')}")
        text_parts.append(f"Date: {msg.get('Date', 'N/A')}")
        text_parts.append("")

        # Body
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                if content_type == 'text/plain':
                    payload = part.get_payload(decode=True)
                    if payload:
                        try:
                            text_parts.append(payload.decode('utf-8'))
                        except UnicodeDecodeError:
                            text_parts.append(payload.decode('latin-1'))
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                try:
                    text_parts.append(payload.decode('utf-8'))
                except UnicodeDecodeError:
                    text_parts.append(payload.decode('latin-1'))

        return '\n'.join(text_parts)


file_processor = FileProcessor()
